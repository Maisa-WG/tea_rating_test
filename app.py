import streamlit as st
import os
import json
import requests
import numpy as np
import faiss
import time
import pickle
from github import Github, GithubException, Auth  # 新增 Auth
from pathlib import Path
from io import BytesIO
from typing import List, Dict, Any, Tuple, Optional
from PyPDF2 import PdfReader
from http import HTTPStatus
import dashscope
from dashscope import TextEmbedding
from openai import OpenAI
from docx import Document
import matplotlib.pyplot as plt
from scipy.interpolate import make_interp_spline
import base64

# ==========================================
# [SECTION 0] 基础配置与路径定义
# ==========================================

st.set_page_config(
    page_title="茶饮六因子AI评分器 Pro",
    page_icon="🍵",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 样式定义
st.markdown("""
    <style>
    .main-title {font-size: 2.5em; font-weight: bold; text-align: center; color: #2E7D32; margin-bottom: 0.5em;}
    .slogan {font-size: 1.2em; font-style: italic; text-align: center; color: #558B2F; margin-bottom: 30px; font-family: "KaiTi", "楷体", serif;}
    .factor-card {background-color: #F1F8E9; padding: 15px; border-radius: 10px; margin-bottom: 10px; border-left: 5px solid #4CAF50;}
    .score-header {display:flex; justify-content:space-between; font-weight:bold; color:#2E7D32;}
    .advice-tag {font-size: 0.85em; padding: 2px 6px; border-radius: 4px; margin-top: 5px; background-color: #fff; border: 1px dashed #4CAF50; color: #388E3C; display: inline-block;}
    .master-comment {background-color: #FFFDE7; border: 1px solid #FFF9C4; padding: 15px; border-radius: 8px; font-family: "KaiTi", serif; font-size: 1.1em; color: #5D4037; margin-bottom: 20px; line-height: 1.6;}
    .ft-card {border: 1px solid #ddd; padding: 15px; border-radius: 8px; background-color: #f8f9fa; margin-top: 10px;}
    .case-card {border: 1px solid #e0e0e0; padding: 12px; border-radius: 8px; margin-bottom: 10px; background-color: #fafafa;}
    </style>
""", unsafe_allow_html=True)

class PathConfig:
    """路径管理类"""
    # 外部资源文件（位于同级目录）
    SRC_SYS_PROMPT = Path("sys_p.txt")
    SRC_CASES = Path("tea_data/case.json")  # Case文件存储
    # 运行时数据目录
    DATA_DIR = Path("./tea_data")
    RAG_DIR = Path("./tea_data/RAG")  # RAG文件存储目录
    
    def __init__(self):
        self.DATA_DIR.mkdir(exist_ok=True)
        self.RAG_DIR.mkdir(exist_ok=True)  # 确保RAG目录存在
        # 向量库与持久化数据
        self.kb_index = self.DATA_DIR / "kb.index"
        self.kb_chunks = self.DATA_DIR / "kb_chunks.pkl"
        self.kb_files = self.DATA_DIR / "kb_files.json"  # 新增：记录RAG文件列表
        self.case_index = self.DATA_DIR / "cases.index"
        self.case_data = self.DATA_DIR / "case.json"  # 修改：与GitHub保持一致
        
        # 微调与Prompt配置
        self.training_file = self.DATA_DIR / "deepseek_finetune.jsonl"
        self.ft_status = self.DATA_DIR / "ft_status.json"
        self.prompt_config_file = self.DATA_DIR / "prompts.json"

PATHS = PathConfig()

# 默认的用户Prompt模板（System Prompt将从文件读取）
DEFAULT_USER_TEMPLATE = """【待评分产品】
{product_desc}

【参考标准（知识库）】
{context_text}

【相似判例得分参考（案例库）】
{case_text}

请严格输出以下JSON格式（不含Markdown）：
{{
  "master_comment": "约100字的宗师级总评，富含文化意蕴...",
  "scores": {{
    "优雅性": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "辨识度": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "协调性": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "饱和度": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "持久性": {{"score": 0-9, "comment": "...", "suggestion": "..."}},
    "苦涩度": {{"score": 0-9, "comment": "...", "suggestion": "..."}}
  }}
}}"""

# ==========================================
# [SECTION 1] 资源与数据管理
# ==========================================

class ResourceManager:
    """负责外部文件加载、数据持久化及格式转换"""

    @staticmethod
    def load_external_text(path: Path, fallback: str = "") -> str:
        """读取外部文本文件 (如 sys_p.txt)"""
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return f.read().strip()
            except Exception as e:
                st.error(f"加载文件 {path} 失败: {e}")
        return fallback

    @staticmethod
    def load_external_json(path: Path, fallback: Any = None) -> Any:
        """读取外部JSON文件"""
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception as e:
                st.error(f"加载文件 {path} 失败: {e}")
        return fallback if fallback is not None else []

    @staticmethod
    def save(index: Any, data: Any, idx_path: Path, data_path: Path, is_json: bool = False):
        """保存 FAISS 索引和数据文件"""
        if index: faiss.write_index(index, str(idx_path))
        with open(data_path, "w" if is_json else "wb", encoding="utf-8" if is_json else None) as f:
            if is_json: json.dump(data, f, ensure_ascii=False, indent=2)
            else: pickle.dump(data, f)
    
    @staticmethod
    def load(idx_path: Path, data_path: Path, is_json: bool = False) -> Tuple[Any, List]:
        """加载 FAISS 索引和数据文件"""
        if idx_path.exists() and data_path.exists():
            try:
                index = faiss.read_index(str(idx_path))
                with open(data_path, "r" if is_json else "rb", encoding="utf-8" if is_json else None) as f:
                    data = json.load(f) if is_json else pickle.load(f)
                return index, data
            except: pass
        return faiss.IndexFlatIP(1024), []

    # ===== 微调相关方法 =====
    @staticmethod
    def overwrite_finetune(cases: List[Dict], sys_prompt: str, user_tpl: str) -> int:
        """覆盖写入微调数据集 (.jsonl) - 修改为覆盖逻辑"""
        try:
            count = 0
            with open(PATHS.training_file, "w", encoding="utf-8") as f:
                for c in cases:
                    case_text = c.get("text", "")
                    scores = c.get("scores", {})
                    master_comment = c.get("master_comment", "（人工校准）")
                    
                    user_content = user_tpl.format(product_desc=case_text, context_text="", case_text="")
                    assistant_content = json.dumps({"master_comment": master_comment, "scores": scores}, ensure_ascii=False)
                    entry = {
                        "messages": [
                            {"role": "system", "content": sys_prompt},
                            {"role": "user", "content": user_content},
                            {"role": "assistant", "content": assistant_content}
                        ]
                    }
                    f.write(json.dumps(entry, ensure_ascii=False) + "\n")
                    count += 1
            return count
        except Exception as e:
            print(f"[ERROR] Finetune overwrite failed: {e}")
            return 0

    @staticmethod
    def save_ft_status(job_id, status, fine_tuned_model=None):
        data = {"job_id": job_id, "status": status, "timestamp": time.time()}
        if fine_tuned_model: data["fine_tuned_model"] = fine_tuned_model
        with open(PATHS.ft_status, 'w') as f: json.dump(data, f)

    @staticmethod
    def load_ft_status():
        if PATHS.ft_status.exists():
            try: return json.load(open(PATHS.ft_status, 'r'))
            except: pass
        return None

    # ===== RAG文件管理 =====
    @staticmethod
    def save_kb_files(file_list: List[str]):
        """保存知识库文件列表"""
        with open(PATHS.kb_files, "w", encoding="utf-8") as f:
            json.dump(file_list, f, ensure_ascii=False, indent=2)
    
    @staticmethod
    def load_kb_files() -> List[str]:
        """加载知识库文件列表"""
        if PATHS.kb_files.exists():
            try:
                with open(PATHS.kb_files, "r", encoding="utf-8") as f:
                    return json.load(f)
            except: pass
        return []

# ==========================================
# [SECTION 1.5] Github 同步工具 (增强版)
# ==========================================

class GithubSync:
    """负责将数据同步回 Github 仓库"""
    
    @staticmethod
    def _get_github_config():
        """获取GitHub配置"""
        token = st.secrets.get("GITHUB_TOKEN")
        repo_name = st.secrets.get("GITHUB_REPO")
        branch = st.secrets.get("GITHUB_BRANCH", "main")
        return token, repo_name, branch
    
    @staticmethod
    def _get_github_client():
        """获取 GitHub 客户端（使用新的认证方式）"""
        token, repo_name, branch = GithubSync._get_github_config()
        if not token or not repo_name:
            return None, None, None
        # 使用新的认证方式，避免 DeprecationWarning
        g = Github(auth=Auth.Token(token))
        return g, repo_name, branch
    
    @staticmethod
    def push_json(file_path_in_repo: str, data_dict: Dict, commit_msg: str = "Update via Streamlit") -> bool:
        """推送 JSON 数据到 Github"""
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            st.error("❌ 未配置 Github Token 或 仓库名 (GITHUB_TOKEN / GITHUB_REPO)")
            return False

        try:
            repo = g.get_repo(repo_name)
            content_str = json.dumps(data_dict, ensure_ascii=False, indent=2)
            
            try:
                contents = repo.get_contents(file_path_in_repo, ref=branch)
                repo.update_file(
                    path=contents.path,
                    message=commit_msg,
                    content=content_str,
                    sha=contents.sha,
                    branch=branch
                )
            except GithubException as e:
                if e.status == 404:
                    repo.create_file(
                        path=file_path_in_repo,
                        message=f"Create {file_path_in_repo}",
                        content=content_str,
                        branch=branch
                    )
                else:
                    raise e
            return True

        except Exception as e:
            st.error(f"Github 同步失败: {str(e)}")
            return False
    @staticmethod
    def load_json(file_path_in_repo: str, default=None):
        """从 Github 读取 JSON 文件；不存在/读取失败则返回 default"""
        if default is None:
            default = []
    
        g, repo_name, branch = GithubSync._get_github_client()
        if not g or not repo_name:
            return default
    
        try:
            repo = g.get_repo(repo_name)
            contents = repo.get_contents(file_path_in_repo, ref=branch)
            raw = contents.decoded_content.decode("utf-8").strip()
            if not raw:
                return default
            return json.loads(raw)
    
        except GithubException as e:
            if getattr(e, "status", None) == 404:
                return default
            st.error(f"Github 读取失败: {str(e)}")
            return default
    
        except Exception as e:
            st.error(f"Github 读取失败: {str(e)}")
            return default
    
    @staticmethod
    def push_binary_file(file_path_in_repo: str, file_content: bytes, commit_msg: str = "Upload file") -> bool:
            """推送二进制文件到 Github (如PDF, DOCX等)
            
            重要：PyGithub的create_file/update_file接受bytes类型时会自动进行base64编码
            不要手动编码，否则会导致双重编码！
            """
            g, repo_name, branch = GithubSync._get_github_client()
            
            if not g or not repo_name:
                st.error("❌ 未配置 Github Token 或 仓库名")
                return False
    
            try:
                repo = g.get_repo(repo_name)
                # 注意：直接传bytes，PyGithub会自动base64编码
                # 不要手动编码！否则会导致双重编码，文件损坏
                
                try:
                    contents = repo.get_contents(file_path_in_repo, ref=branch)
                    repo.update_file(
                        path=contents.path,
                        message=commit_msg,
                        content=file_content,  # 直接传bytes
                        sha=contents.sha,
                        branch=branch
                    )
                except GithubException as e:
                    if e.status == 404:
                        repo.create_file(
                            path=file_path_in_repo,
                            message=f"Create {file_path_in_repo}",
                            content=file_content,  # 直接传bytes
                            branch=branch
                        )
                    else:
                        raise e
                return True
    
            except Exception as e:
                st.error(f"Github 文件上传失败: {str(e)}")
                return False
    
    @staticmethod
    def delete_file(file_path_in_repo: str, commit_msg: str = "Delete file") -> bool:
        """从 Github 删除文件"""
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            return False

        try:
            repo = g.get_repo(repo_name)
            
            try:
                contents = repo.get_contents(file_path_in_repo, ref=branch)
                repo.delete_file(
                    path=contents.path,
                    message=commit_msg,
                    sha=contents.sha,
                    branch=branch
                )
                return True
            except GithubException as e:
                if e.status == 404:
                    return True  # 文件本来就不存在
                raise e

        except Exception as e:
            st.error(f"Github 删除文件失败: {str(e)}")
            return False

    @staticmethod
    def add_rag_files(uploaded_files: List, rag_folder: str = "tea_data/RAG") -> Tuple[bool, List[str]]:
        """
        添加RAG文件到GitHub（只添加，不删除现有文件）
        - uploaded_files: Streamlit上传的文件对象列表
        - rag_folder: GitHub上的RAG文件夹路径
        返回: (是否成功, 成功上传的文件名列表)
        """
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            st.error("❌ 未配置 Github Token 或 仓库名")
            return False, []

        try:
            uploaded_names = []
            for uf in uploaded_files:
                file_path = f"{rag_folder}/{uf.name}"
                uf.seek(0)
                file_content = uf.read()
                if GithubSync.push_binary_file(file_path, file_content, f"Add RAG file: {uf.name}"):
                    uploaded_names.append(uf.name)
                else:
                    st.warning(f"⚠️ 上传 {uf.name} 失败")
            
            return len(uploaded_names) > 0, uploaded_names

        except Exception as e:
            st.error(f"RAG文件添加失败: {str(e)}")
            return False, []

    @staticmethod
    def list_rag_files(rag_folder: str = "tea_data/RAG") -> List[str]:
        """
        获取GitHub上RAG文件夹中的所有文件名
        返回: 文件名列表
        """
        g, repo_name, branch = GithubSync._get_github_client()
        
        if not g or not repo_name:
            return []

        try:
            repo = g.get_repo(repo_name)
            contents = repo.get_contents(rag_folder, ref=branch)
            return [c.name for c in contents if c.type == "file"]
        except GithubException as e:
            if e.status == 404:
                return []  # 文件夹不存在
            print(f"[ERROR] 获取RAG文件列表失败: {e}")
            return []
        except Exception as e:
            print(f"[ERROR] 获取RAG文件列表失败: {e}")
            return []

    @staticmethod
    def delete_rag_file(filename: str, rag_folder: str = "tea_data/RAG") -> bool:
        """
        从GitHub删除单个RAG文件
        - filename: 要删除的文件名
        - rag_folder: GitHub上的RAG文件夹路径
        返回: 是否成功
        """
        file_path = f"{rag_folder}/{filename}"
        return GithubSync.delete_file(file_path, f"Delete RAG file: {filename}")

    @staticmethod
    def sync_cases(cases: List[Dict], file_path: str = "tea_data/case.json") -> bool:
        """同步判例库到GitHub"""
        return GithubSync.push_json(file_path, cases, "Update case.json from App")

    @staticmethod
    def pull_rag_folder(rag_folder: str = "tea_data/RAG") -> List[Tuple[str, bytes]]:
        """
        从 GitHub 拉取 RAG 文件夹中的所有文件
        返回: [(文件名, 文件内容bytes), ...]
        
        优化策略：
        1. 优先使用 Raw URL（最可靠，适合大文件）
        2. 备用方案：Git Blob API（仅小于1MB的文件）
        3. 增加完整性验证：对比文件大小
        4. 支持重试机制
        """
        token, repo_name, branch = GithubSync._get_github_config()
        
        if not token or not repo_name:
            print("[WARN] GitHub config not found, skip pulling RAG")
            return []

        def download_with_retry(url, headers, max_retries=3):
            """带重试的下载函数"""
            for attempt in range(1, max_retries + 1):
                try:
                    response = requests.get(url, headers=headers, timeout=180, stream=True)
                    if response.status_code == 200:
                        # 使用 stream=True 分块下载，避免大文件超时
                        content = b''
                        for chunk in response.iter_content(chunk_size=8192):
                            if chunk:
                                content += chunk
                        return content, True
                    else:
                        print(f"[WARN]     尝试 {attempt}/{max_retries}: HTTP {response.status_code}")
                except Exception as e:
                    print(f"[WARN]     尝试 {attempt}/{max_retries}: {e}")
                    if attempt < max_retries:
                        import time
                        time.sleep(2)  # 等待2秒后重试
            return None, False

        try:
            g = Github(auth=Auth.Token(token))
            repo = g.get_repo(repo_name)
            
            files = []
            print(f"[INFO] ========== 开始从 GitHub 拉取 RAG 文件 ==========")
            print(f"[INFO] 仓库: {repo_name}, 分支: {branch}, 文件夹: {rag_folder}")
            
            try:
                contents = repo.get_contents(rag_folder, ref=branch)
                file_list = [c for c in contents if c.type == "file"]
                print(f"[INFO] 发现 {len(file_list)} 个文件")
                
                for idx, content in enumerate(file_list, 1):
                    print(f"\n[INFO] [{idx}/{len(file_list)}] 正在处理: {content.name}")
                    print(f"[INFO]   → 期望大小: {content.size:,} bytes")
                    file_content = None
                    download_method = None
                    
                    # ===== 方法1：Raw URL（优先，最可靠） =====
                    print(f"[INFO]   → 方法1: Raw URL 下载...")
                    raw_url = f"https://raw.githubusercontent.com/{repo_name}/{branch}/{rag_folder}/{content.name}"
                    headers = {"Authorization": f"Bearer {token}"}
                    file_content, success = download_with_retry(raw_url, headers, max_retries=3)
                    
                    if success and file_content:
                        download_method = "Raw URL"
                        print(f"[INFO]   ✓ 下载完成: {len(file_content):,} bytes")
                    
                    # ===== 方法2：Git Blob API（仅用于小文件 <1MB） =====
                    if file_content is None and content.size < 1024 * 1024:  # 1MB
                        try:
                            print(f"[INFO]   → 方法2: Git Blob API...")
                            blob = repo.get_git_blob(content.sha)
                            if blob.encoding == "base64":
                                file_content = base64.b64decode(blob.content)
                                download_method = "Git Blob"
                                print(f"[INFO]   ✓ 下载完成: {len(file_content):,} bytes")
                        except Exception as e:
                            print(f"[WARN]   ✗ Git Blob 失败: {e}")
                    
                    # ===== 方法3：Download URL（兜底） =====
                    if file_content is None and content.download_url:
                        print(f"[INFO]   → 方法3: Download URL...")
                        headers = {
                            "Authorization": f"Bearer {token}",
                            "Accept": "application/vnd.github.v3.raw"
                        }
                        file_content, success = download_with_retry(content.download_url, headers, max_retries=3)
                        if success:
                            download_method = "Download URL"
                            print(f"[INFO]   ✓ 下载完成: {len(file_content):,} bytes")
                    
                    # ===== 验证文件完整性 =====
                    if file_content:
                        actual_size = len(file_content)
                        expected_size = content.size
                        
                        print(f"[INFO]   → 验证完整性...")
                        print(f"[INFO]     期望: {expected_size:,} bytes")
                        print(f"[INFO]     实际: {actual_size:,} bytes")
                        
                        if actual_size == expected_size:
                            files.append((content.name, file_content))
                            print(f"[INFO]   ✅ {content.name} 完整性验证通过 (方法: {download_method})")
                        else:
                            size_diff = abs(actual_size - expected_size)
                            print(f"[ERROR]  ❌ {content.name} 大小不匹配 (差异: {size_diff:,} bytes)")
                            print(f"[ERROR]     文件可能损坏，跳过...")
                    else:
                        print(f"[ERROR]  ❌ {content.name} 所有下载方法均失败")
                            
            except GithubException as e:
                if e.status == 404:
                    print(f"[INFO] RAG 文件夹不存在: {rag_folder}")
                    return []
                print(f"[ERROR] GitHub API 异常: {e}")
                raise e
            
            print(f"\n[INFO] ========== RAG 拉取完成: {len(files)}/{len(file_list)} 个文件验证通过 ==========\n")
            return files

        except Exception as e:
            print(f"[ERROR] 拉取 RAG 文件夹失败: {e}")
            import traceback
            traceback.print_exc()
            return []

# --- [新增] 日志与评测管理类 ---
class EvaluationLogger:
    FILE_NAME = "tea_data/eval_logs.json"

    @staticmethod
    def load_logs():
        """从 GitHub 同步并加载日志"""
        content = GithubSync.load_json(EvaluationLogger.FILE_NAME)
        return content if content else []

    @staticmethod
    def log_evaluation(text, model_output, expert_output, model_name="Qwen2.5-7B-Instruct"):
        """
        核心：同时记录 AI 的原始输出和专家的校准结果
        """
        logs = EvaluationLogger.load_logs()
        new_entry = {
            "id": str(int(time.time())),
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
            "input_text": text,
            "model_prediction": model_output, # 原始预测包 (scores + master_comment)
            "expert_ground_truth": expert_output, # 专家修正包 (scores + master_comment)
            "analysis": None,
            "meta": {"model": model_name}
        }
        logs.insert(0, new_entry) # 最新记录在前
        if len(logs) > 500: logs = logs[:500] # 限制日志长度
        GithubSync.push_json(EvaluationLogger.FILE_NAME, logs, f"Eval log {new_entry['id']}")
        return logs

    @staticmethod
    def run_judge(log_id, llm_client):
        """运行 LLM 裁判：分析 AI 为什么评错了"""
        logs = EvaluationLogger.load_logs()
        target = next((l for l in logs if l["id"] == log_id), None)
        if not target or not target.get("expert_ground_truth"): return "缺少对比数据"

        judge_prompt = f"""
        你是一名茶叶感官审评专家教练。请对比以下“模型评分”与“专家标准评分”，分析差异原因。
        【原始评语】: {target['input_text']}
        【模型原始分】: {json.dumps(target['model_prediction'], ensure_ascii=False)}
        【专家校准分】: {json.dumps(target['expert_ground_truth'], ensure_ascii=False)}
        请输出简短的误差分析：
        1. 哪些维度偏差较大？2. 模型误解了评语中的哪个关键描述？3. 针对此案例，应如何优化 Prompt？
        """
        try:
            resp = llm_client.chat.completions.create(
                model="deepseek-chat",
                messages=[{"role": "user", "content": judge_prompt}]
            )
            analysis = resp.choices[0].message.content
            # 更新日志并同步
            for l in logs:
                if l["id"] == log_id: l["analysis"] = analysis
            GithubSync.push_json(EvaluationLogger.FILE_NAME, logs, f"Update judge {log_id}")
            return analysis
        except Exception as e:
            return f"裁判分析失败: {str(e)}"
# ==========================================
# [SECTION 2] AI 服务 (Embedding & LLM)
# ==========================================

class AliyunEmbedder:
    def __init__(self, api_key):
        self.model_name = "text-embedding-v3"
        dashscope.api_key = api_key # 确保 API KEY 被正确设置给全局
        
    def encode(self, texts: List[str]) -> np.ndarray:
        if not texts:
            return np.zeros((0, 1024), dtype="float32")
        if isinstance(texts, str):
            texts = [texts]
    
        try:
            resp = TextEmbedding.call(model=self.model_name, input=texts)
        except Exception as e:
            # 关键：不要吞掉，否则检索会“永远不变”
            raise RuntimeError(f"[Embedding] call failed: {type(e).__name__}: {e}")
    
        if resp.status_code != HTTPStatus.OK:
            # 把错误信息暴露出来
            msg = getattr(resp, "message", "")
            raise RuntimeError(f"[Embedding] HTTP not OK: {resp.status_code}, message={msg}")
    
        vecs = np.array([i["embedding"] for i in resp.output["embeddings"]], dtype="float32")
    
        # 关键：检测全 0 / 常量向量（典型“检索永远一样”）
        norms = np.linalg.norm(vecs, axis=1)
        if np.any(norms < 1e-6):
            raise RuntimeError("[Embedding] got near-zero vector(s). Check ALIYUN_API_KEY / dashscope / model availability.")
    
        return vecs


def _ensure_ip_index_from_texts(texts, embedder):
    """Build an IndexFlatIP (cosine via L2-normalized vectors) from texts."""
    if not texts:
        return faiss.IndexFlatIP(1024)
    vecs = embedder.encode(texts).astype("float32")
    # normalize for cosine similarity
    faiss.normalize_L2(vecs)
    dim = vecs.shape[1] if vecs.ndim == 2 and vecs.shape[0] > 0 else 1024
    idx = faiss.IndexFlatIP(dim)
    if vecs.shape[0] > 0:
        idx.add(vecs)
    return idx

def ensure_case_index_cosine(embedder):
    """Migrate / rebuild case index to cosine (IP + normalized vectors) if needed."""
    case_idx, case_data = st.session_state.get("cases", (faiss.IndexFlatIP(1024), []))
    # If empty, nothing to do
    if not case_data:
        st.session_state.cases = (faiss.IndexFlatIP(1024), case_data)
        return

    # If metric not IP or ntotal mismatch, rebuild
    metric = getattr(case_idx, "metric_type", None)
    if metric != faiss.METRIC_INNER_PRODUCT or case_idx.ntotal != len(case_data):
        texts = [c.get("text", "") for c in case_data]
        new_idx = _ensure_ip_index_from_texts(texts, embedder)
        st.session_state.cases = (new_idx, case_data)
        ResourceManager.save(new_idx, case_data, PATHS.case_index, PATHS.case_data, is_json=True)

def ensure_kb_index_cosine(embedder):
    """Migrate / rebuild kb index to cosine (IP + normalized vectors) if needed."""
    kb_idx, kb_chunks = st.session_state.get("kb", (faiss.IndexFlatIP(1024), []))
    if not kb_chunks:
        st.session_state.kb = (faiss.IndexFlatIP(1024), kb_chunks)
        return

    metric = getattr(kb_idx, "metric_type", None)
    if metric != faiss.METRIC_INNER_PRODUCT or kb_idx.ntotal != len(kb_chunks):
        new_idx = _ensure_ip_index_from_texts(kb_chunks, embedder)
        st.session_state.kb = (new_idx, kb_chunks)
        ResourceManager.save(new_idx, kb_chunks, PATHS.kb_index, PATHS.kb_chunks)


def llm_normalize_user_input(raw_query: str, client: OpenAI) -> str:
    """使用 LLM 对用户输入做语义规范化 / 去噪"""
    system_prompt = (
        """
          A. 角色与目标
          你是"茶评清洗器"。你的任务是从输入文本中提取并输出只与茶评相关的信息，删除无关内容，保持原意与原有表述风格，只能删减不能修改。
          B. 什么算"相关信息"（保留）
          仅保留与以下内容有关的句子/短语：
          茶的基本信息：茶名/品类、产地、年份、工艺、等级、原料、香型等
          干茶/茶汤/叶底：外观、色泽、条索、汤色、叶底描述
          香气与滋味：香气类型、强弱、层次、回甘、生津、涩感、苦感、甜度、醇厚度、喉韵、体感等
          冲泡信息与表现：器具、投茶量、水温、时间、出汤、几泡变化、耐泡度、适饮建议
          主观评价与结论：好喝/一般/缺点/性价比
          C. 什么算"无关信息"（删除）
          删除与茶评无直接关系的内容，例如：
          与茶无关的生活日常、情绪宣泄、社交聊天、段子
          店铺/物流/客服/包装破损/发货慢（除非"包装异味影响茶"这类直接影响品饮）
          广告、价格链接、优惠券、引流话术、品牌吹水（除非是"性价比"且与品饮结论相关）
          与其它产品/话题无关的对比闲聊
          凑字数内容
          D. 输出格式
          只输出清洗后的茶评正文，不要解释、不加标题、不输出"删除了什么"
          如果输入中没有任何茶评相关信息，则输出："无相关茶评信息"
          E. 操作原则
          尽量保留原句；只做删除/少量拼接
          不要补充不存在的细节，不要推测        
          """
    )

    resp = client.chat.completions.create(
        model="deepseek-chat",
        temperature=0,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": raw_query}
        ]
    )
    return resp.choices[0].message.content.strip()

def run_scoring(text: str, kb_res: Tuple, case_res: Tuple, prompt_cfg: Dict, embedder: AliyunEmbedder, client: OpenAI, model_id: str, k_num: int, c_num: int):
    """执行 RAG 检索与 LLM 评分"""
    vec = embedder.encode([text]).astype("float32")
    faiss.normalize_L2(vec)
    # Debug：norm 应该约等于 1.0；如果不是，embedding 或 normalize 有问题
    print(f"[DEBUG] query_vec_norm={float(np.linalg.norm(vec[0])):.6f}")

    # --- KB ---
    ctx_txt, hits = "（无手册资料）", []
    if kb_res[0].ntotal > 0:
        _, idx = kb_res[0].search(vec, k_num)
        hits = [kb_res[1][i] for i in idx[0] if i < len(kb_res[1])]
        ctx_txt = "\n".join([f"- {h[:200]}..." for h in hits])

    # --- CASES ---
    case_txt, found_cases = "", []
    if case_res[0].ntotal > 0:
        _, idx = case_res[0].search(vec, c_num)
        for i in idx[0]:
            if 0 <= i < len(case_res[1]):
                c = case_res[1][i]
                found_cases.append(c)

                score_details = []
                for factor, info in c.get("scores", {}).items():
                    if isinstance(info, dict):
                        score_details.append(
                            f"{factor}: {info.get('score')}分 (理由: {info.get('comment', '无')})"
                        )
                scores_str = " | ".join(score_details)

                case_txt += (
                    f"\n---\n"
                    f"【相似判例】: {c.get('text','')}\n"
                    f"【该判例专家分】{scores_str}\n"
                    f"【硬约束】如果待评分文本与该判例高度一致（语义基本相同），六因子分数应优先对齐该判例的专家分；只有明确出现相反描述时才允许偏离，并必须在comment里解释偏离原因。\n"
                )

    if not found_cases:
        case_txt = "（无相似判例）"

    sys_p = prompt_cfg.get('system_template', "")
    user_p = prompt_cfg.get('user_template', "").format(product_desc=text, context_text=ctx_txt, case_text=case_txt)

    try:
        resp = client.chat.completions.create(
            model=model_id,
            messages=[{"role":"system", "content":sys_p}, {"role":"user", "content":user_p}],
            response_format={"type": "json_object"},
            temperature=0.3
        )
        return json.loads(resp.choices[0].message.content), hits, found_cases
    except Exception as e:
        st.error(f"Inference Error: {e}")
        return None, [], []

# ==========================================
# [SECTION 3] 辅助与可视化
# ==========================================

def parse_file(uploaded_file) -> str:
    """解析上传文件"""
    try:
        if uploaded_file.name.endswith('.txt'): return uploaded_file.read().decode("utf-8")
        if uploaded_file.name.endswith('.pdf'): return "".join([p.extract_text() for p in PdfReader(uploaded_file).pages])
        if uploaded_file.name.endswith('.docx'): return "\n".join([p.text for p in Document(uploaded_file).paragraphs])
    except: return ""
    return ""

def parse_file_bytes(filename: str, content: bytes) -> str:
    """
    解析文件内容 (从 bytes) - 用于从 GitHub 拉取的文件
    支持格式: .txt, .pdf, .docx
    """
    try:
        # 1. 处理 TXT 文件
        if filename.lower().endswith('.txt'):
            text = content.decode('utf-8', errors='ignore')
            print(f"[INFO]     → TXT 解析成功: {len(text)} 字符")
            return text
        
        # 2. 处理 PDF 文件
        elif filename.lower().endswith('.pdf'):
            try:
                print(f"[INFO]     → 开始解析 PDF...")
                print(f"[INFO]     → 文件大小: {len(content):,} bytes")
                
                # 验证 PDF 文件头
                if not content.startswith(b'%PDF'):
                    print(f"[ERROR]    → 不是有效的 PDF 文件（文件头错误）")
                    print(f"[ERROR]    → 前20字节: {content[:20]}")
                    return ""
                
                # 验证 PDF 文件尾
                if b'%%EOF' not in content[-1024:]:
                    print(f"[WARN]     → PDF 文件尾标记缺失，文件可能不完整")
                    print(f"[WARN]     → 后50字节: {content[-50:]}")
                
                # 尝试解析 PDF
                reader = PdfReader(BytesIO(content))
                page_count = len(reader.pages)
                print(f"[INFO]     → PDF 共 {page_count} 页")
                
                text = ""
                failed_pages = []
                
                for idx, page in enumerate(reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            if idx % 10 == 0:  # 每 10 页输出一次进度
                                print(f"[INFO]     → 已处理 {idx}/{page_count} 页")
                    except Exception as e:
                        failed_pages.append(idx)
                        print(f"[WARN]     → 第 {idx} 页解析失败: {e}")
                        continue
                
                if failed_pages:
                    print(f"[WARN]     → 共 {len(failed_pages)} 页解析失败: {failed_pages[:10]}{'...' if len(failed_pages) > 10 else ''}")
                
                if text.strip():
                    print(f"[INFO]     → PDF 解析完成: {len(text):,} 字符 (成功率: {(page_count-len(failed_pages))/page_count*100:.1f}%)")
                    return text
                else:
                    print(f"[WARN]     → PDF 解析结果为空")
                    return ""
                    
            except Exception as e:
                print(f"[ERROR]    ✗ PDF 解析失败: {type(e).__name__}: {e}")
                
                # 如果是 EOF 错误，提供更多诊断信息
                if "EOF" in str(e) or "PdfReadError" in str(type(e).__name__):
                    print(f"[ERROR]    → 这通常意味着 PDF 文件下载不完整或损坏")
                    print(f"[ERROR]    → 文件大小: {len(content):,} bytes")
                    print(f"[ERROR]    → 文件头: {content[:20]}")
                    print(f"[ERROR]    → 文件尾: {content[-50:] if len(content) > 50 else content}")
                
                import traceback
                traceback.print_exc()
                return ""
        
        # 3. 处理 DOCX 文件
        elif filename.lower().endswith('.docx'):
            doc = Document(BytesIO(content))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            print(f"[INFO]     → DOCX 解析成功: {len(text)} 字符")
            return text
        
        else:
            print(f"[WARN]     → 不支持的文件格式: {filename}")
            return ""
            
    except Exception as e:
        print(f"[ERROR]    ✗ 解析 {filename} 失败: {e}")
        import traceback
        traceback.print_exc()
        return ""

def create_word_report(results: List[Dict]) -> BytesIO:
    """生成Word报告"""
    doc = Document()
    doc.add_heading("茶评批量评分报告", 0)
    for item in results:
        doc.add_heading(f"条目 {item['id']}", 1)
        doc.add_paragraph(f"原文：{item['text']}")
        s = item.get('scores', {}).get('scores', {})
        mc = item.get('scores', {}).get('master_comment', '')
        if mc: doc.add_paragraph(f"总评：{mc}", style="Intense Quote")
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = '因子', '分数', '评语', '建议'
        for k, v in s.items():
            r = table.add_row().cells
            r[0].text = k
            r[1].text = str(v.get('score',''))
            r[2].text = v.get('comment','')
            r[3].text = v.get('suggestion','')
        doc.add_paragraph("_"*20)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def plot_flavor_shape(scores_data: Dict):
    """绘制风味形态图"""
    s = scores_data["scores"]
    top = (s["优雅性"]["score"] + s["辨识度"]["score"]) / 2
    mid = (s["协调性"]["score"] + s["饱和度"]["score"]) / 2
    base = (s["持久性"]["score"] + s["苦涩度"]["score"]) / 2
    
    fig, ax = plt.subplots(figsize=(4, 5))
    fig.patch.set_alpha(0); ax.patch.set_alpha(0)

    y = np.array([1, 2, 3]) 
    x = np.array([base, mid, top])
    y_new = np.linspace(1, 3, 300)
    try:
        spl = make_interp_spline(y, x, k=2)
        x_smooth = spl(y_new)
    except:
        x_smooth = np.interp(y_new, y, x)
    x_smooth = np.maximum(x_smooth, 0.1)

    colors = {'base': '#8B4513', 'mid': '#D2691E', 'top': '#FFD700'}
    for mask, col in [((y_new>=1.0)&(y_new<=1.6), colors['base']), 
                      ((y_new>1.6)&(y_new<=2.4), colors['mid']), 
                      ((y_new>2.4)&(y_new<=3.0), colors['top'])]:
        ax.fill_betweenx(y_new[mask], -x_smooth[mask], x_smooth[mask], color=col, alpha=0.9, edgecolor=None)

    ax.plot(x_smooth, y_new, 'k', linewidth=1, alpha=0.2)
    ax.plot(-x_smooth, y_new, 'k', linewidth=1, alpha=0.2)
    ax.axhline(y=1.6, color='w', linestyle=':', alpha=0.5)
    ax.axhline(y=2.4, color='w', linestyle=':', alpha=0.5)
    
    font = {'ha': 'center', 'va': 'center', 'color': 'white', 'fontweight': 'bold', 'fontsize': 12}
    ax.text(0, 2.7, f"Top\n{top:.1f}", **font)
    ax.text(0, 2.0, f"Mid\n{mid:.1f}", **font)
    ax.text(0, 1.3, f"Base\n{base:.1f}", **font)
    ax.axis('off'); ax.set_xlim(-10, 10); ax.set_ylim(0.8, 3.2)
    return fig

def bootstrap_seed_cases(embedder: AliyunEmbedder):
    """初始化判例库：如果内存/磁盘中为空，则从 case.json 文件读取"""
    case_idx, case_data = st.session_state.cases
    if len(case_data) > 0: return

    # 从外部 JSON 加载 (修改路径)
    seed_cases = ResourceManager.load_external_json(PATHS.SRC_CASES)
    if not seed_cases:
        # 兼容旧路径
        old_path = Path("seed_case.json")
        seed_cases = ResourceManager.load_external_json(old_path)
    
    if not seed_cases:
        st.warning("case.json 未找到或为空，判例库初始化跳过。")
        return

    texts = [c["text"] for c in seed_cases]
    vecs = embedder.encode(texts)
    faiss.normalize_L2(vecs)

    if case_idx.ntotal == 0: case_idx = faiss.IndexFlatIP(1024)
    if len(vecs) > 0:
        case_idx.add(vecs)
        case_data.extend(seed_cases)
        st.session_state.cases = (case_idx, case_data)
        ResourceManager.save(case_idx, case_data, PATHS.case_index, PATHS.case_data, is_json=True)

def load_rag_from_github(aliyun_key: str) -> Tuple[bool, str]:
    """
    从 GitHub 加载 RAG 文件
    返回: (是否成功, 消息)
    """
    print("\n[INFO] ========== 开始从 GitHub 加载 RAG 数据 ==========")
    
    try:
        # 1. 拉取文件
        print("[INFO] 步骤 1/4: 从 GitHub 拉取 RAG 文件...")
        rag_files = GithubSync.pull_rag_folder("tea_data/RAG")
        
        if not rag_files:
            msg = "GitHub 上没有找到 RAG 文件，或所有文件下载失败"
            print(f"[WARN] {msg}")
            return False, msg
        
        print(f"[INFO] 成功拉取并验证 {len(rag_files)} 个文件")
        
        # 2. 解析文件内容
        print("[INFO] 步骤 2/4: 解析文件内容...")
        all_text = ""
        file_names = []
        parse_success = 0
        parse_failed = []
        
        for fname, fcontent in rag_files:
            file_names.append(fname)
            print(f"\n[INFO]   → 解析 {fname} ({len(fcontent):,} bytes)...")
            
            parsed_text = parse_file_bytes(fname, fcontent)
            if parsed_text and len(parsed_text.strip()) > 100:  # 至少要有100个字符
                all_text += parsed_text + "\n"
                parse_success += 1
                print(f"[INFO]   ✅ 成功提取 {len(parsed_text):,} 字符")
            else:
                parse_failed.append(fname)
                print(f"[WARN]   ❌ 提取失败或文本过短 ({len(parsed_text) if parsed_text else 0} 字符)")
        
        print(f"\n[INFO] 文件解析完成: {parse_success}/{len(rag_files)} 成功")
        if parse_failed:
            print(f"[WARN] 解析失败的文件: {', '.join(parse_failed)}")
        
        if not all_text.strip():
            msg = f"无法从 RAG 文件中提取有效文本（共尝试 {len(rag_files)} 个文件）"
            print(f"[ERROR] {msg}")
            return False, msg
        
        # 3. 切片
        print(f"\n[INFO] 步骤 3/4: 将文本切片...")
        print(f"[INFO]   → 总文本长度: {len(all_text):,} 字符")
        chunks = [all_text[i:i+600] for i in range(0, len(all_text), 500)]
        print(f"[INFO]   ✓ 切片完成: {len(chunks)} 个片段")
        
        if not chunks:
            msg = "切片失败"
            print(f"[ERROR] {msg}")
            return False, msg
        
        # 4. 向量化并构建索引
        print("\n[INFO] 步骤 4/4: 向量化并构建 FAISS 索引...")
        temp_embedder = AliyunEmbedder(aliyun_key)
        kb_idx = faiss.IndexFlatIP(1024)
        
        print(f"[INFO]   → 调用阿里云 Embedding API (批次大小: 25)...")
        
        # 分批向量化，避免 API 限制
        batch_size = 25
        all_vecs = []
        for i in range(0, len(chunks), batch_size):
            batch = chunks[i:i+batch_size]
            try:
                vecs = temp_embedder.encode(batch)
                all_vecs.append(vecs)
                print(f"[INFO]   → 已处理 {min(i+batch_size, len(chunks))}/{len(chunks)} 片段")
            except Exception as e:
                print(f"[WARN]   → 批次 {i}-{i+batch_size} 向量化失败: {e}")
        
        if not all_vecs:
            msg = "向量化失败"
            print(f"[ERROR] {msg}")
            return False, msg
        
        vecs = np.vstack(all_vecs)
        print(f"[INFO]   ✓ 获得向量: {vecs.shape}")
        faiss.normalize_L2(vecs)
        
        kb_idx.add(vecs)
        print(f"[INFO]   ✓ FAISS 索引构建完成 (共 {kb_idx.ntotal} 条)")
        
        # 5. 保存到 session_state 和磁盘
        st.session_state.kb = (kb_idx, chunks)
        st.session_state.kb_files = file_names
        
        ResourceManager.save(kb_idx, chunks, PATHS.kb_index, PATHS.kb_chunks)
        ResourceManager.save_kb_files(file_names)
        
        success_files = [f for f in file_names if f not in parse_failed]
        msg = f"✅ 成功加载 {len(chunks)} 条知识片段\n📁 来源文件: {', '.join(success_files)}"
        if parse_failed:
            msg += f"\n⚠️  解析失败: {', '.join(parse_failed)}"
        
        print(f"[INFO] {msg}")
        print("[INFO] ========== RAG 加载完成 ==========\n")
        return True, msg
        
    except Exception as e:
        msg = f"加载失败: {str(e)}"
        print(f"[ERROR] ❌ {msg}")
        import traceback
        traceback.print_exc()
        print("[INFO] ========== RAG 加载失败 ==========\n")
        return False, msg


# ==========================================
# [SECTION 3.5] 判例管理弹窗
# ==========================================

@st.dialog("📋 判例库管理", width="large")
def show_cases_dialog(embedder: AliyunEmbedder):
    """展示并管理所有判例的弹窗"""
    cases = st.session_state.cases[1]
    
    if not cases:
        st.info("当前判例库为空")
        return
    
    st.write(f"共 **{len(cases)}** 条判例")
    st.caption("💡 勾选要删除的判例，然后点击底部的确认按钮")
    
    # 用于追踪编辑状态
    if 'editing_case_idx' not in st.session_state:
        st.session_state.editing_case_idx = None
    
    # 使用checkbox收集要删除的判例（不会触发rerun导致弹窗关闭）
    selected_to_delete = []
    
    for idx, case in enumerate(cases):
        with st.container(border=True):
            col1, col2, col3 = st.columns([6, 1, 1])
            
            with col1:
                # 显示判例摘要
                text_preview = case.get('text', '')[:100] + ('...' if len(case.get('text', '')) > 100 else '')
                st.markdown(f"**#{idx+1}** {text_preview}")
                
                # 显示分数摘要
                scores = case.get('scores', {})
                if scores:
                    score_str = " | ".join([f"{k}:{v.get('score', '?')}" for k, v in scores.items()])
                    st.caption(score_str)
            
            with col2:
                if st.button("✏️", key=f"edit_{idx}", help="编辑此判例"):
                    st.session_state.editing_case_idx = idx
                    st.rerun()
            
            with col3:
                # 使用checkbox代替button，避免rerun导致弹窗关闭
                if st.checkbox("删除", key=f"del_check_{idx}", label_visibility="collapsed"):
                    selected_to_delete.append(idx)
    
    # 如果有选中要删除的判例
    if selected_to_delete:
        st.warning(f"⚠️ 已选中 {len(selected_to_delete)} 条判例待删除")
        if st.button("✅ 确认删除并同步", type="primary", use_container_width=True):
            # 执行删除
            new_cases = [c for i, c in enumerate(cases) if i not in selected_to_delete]
            
            # 重建FAISS索引
            new_idx = faiss.IndexFlatIP(1024)
            if new_cases:
                texts = [c["text"] for c in new_cases]
                vecs = embedder.encode(texts)
                faiss.normalize_L2(vecs)
                new_idx.add(vecs)
            
            st.session_state.cases = (new_idx, new_cases)
            ResourceManager.save(new_idx, new_cases, PATHS.case_index, PATHS.case_data, is_json=True)
            
            # 同步到GitHub
            with st.spinner("同步到GitHub..."):
                GithubSync.sync_cases(new_cases)
            
            st.success("删除完成！")
            time.sleep(1)
            st.rerun()


@st.dialog("✏️ 编辑判例", width="large")
def edit_case_dialog(case_idx: int, embedder: AliyunEmbedder):
    """编辑单个判例的弹窗"""
    cases = st.session_state.cases[1]
    if case_idx >= len(cases):
        st.error("判例不存在")
        return
    
    case = cases[case_idx]
    factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
    
    st.subheader(f"编辑判例 #{case_idx + 1}")
    
    # 编辑文本
    new_text = st.text_area("判例描述", case.get("text", ""), height=100)
    new_master = st.text_area("总评", case.get("master_comment", ""), height=60)
    new_tags = st.text_input("标签", case.get("tags", ""))
    
    # 编辑各因子分数
    st.markdown("**因子评分**")
    new_scores = {}
    cols = st.columns(3)
    
    old_scores = case.get("scores", {})
    for i, f in enumerate(factors):
        with cols[i % 3]:
            with st.container(border=True):
                st.markdown(f"**{f}**")
                old_f = old_scores.get(f, {})
                new_scores[f] = {
                    "score": st.number_input(f"分数", 0, 9, int(old_f.get("score", 5)), key=f"edit_s_{f}"),
                    "comment": st.text_input(f"评语", old_f.get("comment", ""), key=f"edit_c_{f}"),
                    "suggestion": st.text_input(f"建议", old_f.get("suggestion", ""), key=f"edit_sg_{f}")
                }
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 保存修改并同步", type="primary"):
            # 更新判例
            cases[case_idx] = {
                "text": new_text,
                "scores": new_scores,
                "tags": new_tags,
                "master_comment": new_master,
                "created_at": case.get("created_at", time.strftime("%Y-%m-%d"))
            }
            
            # 重建FAISS索引（因为文本可能变了）
            new_idx = faiss.IndexFlatIP(1024)
            texts = [c["text"] for c in cases]
            vecs = embedder.encode(texts)
            faiss.normalize_L2(vecs)
            new_idx.add(vecs)
            
            st.session_state.cases = (new_idx, cases)
            ResourceManager.save(new_idx, cases, PATHS.case_index, PATHS.case_data, is_json=True)
            
            # 同步到GitHub
            with st.spinner("同步到GitHub..."):
                GithubSync.sync_cases(cases)
            
            st.session_state.editing_case_idx = None
            st.success("保存成功！")
            time.sleep(1)
            st.rerun()
    
    with col2:
        if st.button("❌ 取消"):
            st.session_state.editing_case_idx = None
            st.rerun()


# ==========================================
# [SECTION 4] 主程序逻辑
# ==========================================
# A. 初始化 Session
if 'loaded' not in st.session_state:
    print("\n" + "="*70)
    print("[INFO] ========== 茶饮六因子AI评分器 - 系统初始化 ==========")
    print("="*70)
    
    # 1. 加载本地缓存的 RAG 与判例数据
    print("[INFO] 步骤 1/3: 加载本地缓存数据...")
    kb_idx, kb_data = ResourceManager.load(PATHS.kb_index, PATHS.kb_chunks)
    case_idx, case_data = ResourceManager.load(PATHS.case_index, PATHS.case_data, is_json=True)
    # (removed) case self-heal moved to cosine-migration helper after embedder init

    st.session_state.kb = (kb_idx, kb_data)
    st.session_state.cases = (case_idx, case_data)
    st.session_state.kb_files = ResourceManager.load_kb_files()
    
    print(f"[INFO]   → 知识库: {len(kb_data)} 个片段")
    print(f"[INFO]   → 判例库: {len(case_data)} 条判例")
    print(f"[INFO]   → RAG 文件: {st.session_state.kb_files}")
    
    # 2. 标记是否需要从 GitHub 加载 RAG（延迟加载）
    print("[INFO] 步骤 2/3: 检查 RAG 状态...")
    if len(kb_data) == 0:
        st.session_state.rag_loading_needed = True
        st.session_state.rag_loading_status = "pending"
        print("[INFO]   ⚠️  本地知识库为空，将在应用启动后从 GitHub 加载")
    else:
        st.session_state.rag_loading_needed = False
        st.session_state.rag_loading_status = "complete"
        print(f"[INFO]   ✅ 使用本地缓存: {len(kb_data)} 个片段")
    
    # 3. 加载 Prompt 配置
    print("[INFO] 步骤 3/3: 加载 Prompt 配置...")
    if PATHS.prompt_config_file.exists():
        try:
            with open(PATHS.prompt_config_file, 'r', encoding='utf-8') as f:
                st.session_state.prompt_config = json.load(f)
                print("[INFO]   ✅ 已加载自定义 Prompt 配置")
        except Exception as e:
            print(f"[WARN]   ⚠️  加载失败: {e}，使用默认配置")
        
    if 'prompt_config' not in st.session_state:
        sys_prompt_content = ResourceManager.load_external_text(PATHS.SRC_SYS_PROMPT, fallback="你是一名茶评专家...")
        st.session_state.prompt_config = {
            "system_template": sys_prompt_content,
            "user_template": DEFAULT_USER_TEMPLATE
        }
        print("[INFO]   ✅ 使用默认 Prompt 配置")
    
    st.session_state.loaded = True
    print("="*70)
    print("[INFO] ========== 系统初始化完成（快速启动模式）==========")
    print("="*70 + "\n")



# B. 侧边栏
with st.sidebar:
    st.header("⚙️ 系统配置")
    st.markdown("**🔐 API 配置**")
    aliyun_key = os.getenv("ALIYUN_API_KEY") or st.secrets.get("ALIYUN_API_KEY", "")
    deepseek_key = os.getenv("DEEPSEEK_API_KEY") or st.secrets.get("DEEPSEEK_API_KEY", "")

    if not aliyun_key or not deepseek_key:
        st.warning("⚠️ 未配置 API Key")
        st.stop()
    else:
        st.success("✅ API 就绪")

    st.markdown("---")
    st.markdown(f"**预处理模型：** `Deepseek-chat`")
    st.markdown(f"**评分模型：** `Qwen2.5-7B-Instruct`")
    model_id = "Qwen2.5-7B-Instruct"
    try:
        resp = requests.get("http://117.50.89.74:8001/status", timeout=2)
        if resp.status_code == 200 and resp.json().get("lora_available"):
            model_id = "default_lora"
            st.success("🎉 已启用微调模型")
    except:
        pass
    ft_status = ResourceManager.load_ft_status()
    if ft_status and ft_status.get("status") == "succeeded":
        st.info(f"🎉 发现微调模型：`{ft_status.get('fine_tuned_model')}`")

    embedder = AliyunEmbedder(aliyun_key)
    client = OpenAI(api_key="dummy", base_url="http://117.50.89.74:8000/v1")
    client_d = OpenAI(api_key=deepseek_key, base_url="https://api.deepseek.com")
    
    bootstrap_seed_cases(embedder)
    
    # ✅ Ensure FAISS indices use cosine similarity (IP + normalized vectors)
    ensure_case_index_cosine(embedder)
    ensure_kb_index_cosine(embedder)
    st.markdown("---")
    
    # ===== 延迟加载 RAG 逻辑 =====
    kb_files = st.session_state.get('kb_files', [])
    kb_count = len(st.session_state.kb[1])
    case_count = len(st.session_state.cases[1])
    
    # 检查是否需要从 GitHub 加载 RAG
    if st.session_state.get('rag_loading_needed', False):
        loading_status = st.session_state.get('rag_loading_status', 'pending')
        
        if loading_status == 'pending':
            # 显示加载状态
            with st.status("🔄 正在从 GitHub 加载知识库...", expanded=True) as status:
                st.write("📥 下载 RAG 文件...")
                st.session_state.rag_loading_status = 'loading'
                
                try:
                    # 执行加载
                    success, msg = load_rag_from_github(aliyun_key)
                    
                    if success:
                        status.update(label="✅ 知识库加载完成", state="complete", expanded=False)
                        st.session_state.rag_loading_status = 'complete'
                        st.session_state.rag_loading_needed = False
                        time.sleep(1)
                        st.rerun()
                    else:
                        status.update(label="❌ 知识库加载失败", state="error", expanded=True)
                        st.error(msg)
                        st.info("💡 您可以在 Tab3 手动上传 RAG 文件")
                        st.session_state.rag_loading_status = 'failed'
                        
                        # 添加重试按钮
                        if st.button("🔄 重试加载", type="secondary"):
                            st.session_state.rag_loading_status = 'pending'
                            st.rerun()
                except Exception as e:
                    status.update(label="❌ 加载出错", state="error", expanded=True)
                    st.error(f"加载失败: {str(e)}")
                    st.session_state.rag_loading_status = 'failed'
                    
                    if st.button("🔄 重试加载", type="secondary"):
                        st.session_state.rag_loading_status = 'pending'
                        st.rerun()
        
        elif loading_status == 'loading':
            st.info("🔄 正在加载知识库，请稍候...")
        
        elif loading_status == 'failed':
            st.warning("⚠️ 知识库加载失败")
            if st.button("🔄 重试从 GitHub 加载", type="secondary"):
                st.session_state.rag_loading_status = 'pending'
                st.rerun()
    
    # 更新显示的数据
    kb_count = len(st.session_state.kb[1])
    kb_files = st.session_state.get('kb_files', [])
    
    st.markdown(f"知识库: **{kb_count}** 条 | 判例库: **{case_count}** 条")
    if kb_files:
        pass
    elif kb_count == 0:
        st.caption("⚠️ 知识库为空，请上传文件或从从云端加载")
    
    st.caption("快速上传仅支持.zip文件格式。")
    st.caption("少量文件上传请至\"知识库设计\"板块。")
    
    if st.button("📤 导出数据"):
        import zipfile, shutil
        temp_dir = Path("./temp_export"); temp_dir.mkdir(exist_ok=True)
        for p in [PATHS.kb_index, PATHS.kb_chunks, PATHS.case_index, PATHS.case_data, PATHS.prompt_config_file]:
            if p.exists(): shutil.copy2(p, temp_dir / p.name)
        zip_path = Path("./rag_export.zip")
        with zipfile.ZipFile(zip_path, 'w') as z:
            for f in temp_dir.iterdir(): z.write(f, f.name)
        with open(zip_path, 'rb') as f:
            st.download_button("⬇️ 下载ZIP", f, "tea_data.zip", "application/zip")
        shutil.rmtree(temp_dir); zip_path.unlink()

    if st.button("📥 导入数据"):
        u_zip = st.file_uploader("上传ZIP", type=['zip'])
        if u_zip:
            import zipfile, tempfile
            with tempfile.TemporaryDirectory() as td:
                zp = Path(td)/"u.zip"
                with open(zp,'wb') as f: f.write(u_zip.getvalue())
                with zipfile.ZipFile(zp,'r') as z: z.extractall(PATHS.DATA_DIR)
                st.success("导入成功，请刷新"); st.rerun()

# C. 主界面
st.markdown('<div class="main-title">🍵 茶品六因子 AI 评分器 Pro</div>', unsafe_allow_html=True)
st.markdown('<div class="slogan">"一片叶子落入水中，改变了水的味道..."</div>', unsafe_allow_html=True)

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs(["💡 交互评分", "🚀 批量评分", "📕 知识库设计", "🛠️ 判例库与微调", "📲 提示词（Prompt）配置","测试日志"])

# --- Tab 1: 交互评分 ---
with tab1:
    st.info("将参考知识库与判例库进行评分。确认结果可一键更新判例库。")
    c1, c2, c3, c4, c5 = st.columns([1, 3, 1, 3, 1])
    r_num = c2.number_input("参考知识库条目数量", 1, 20, 3, key="r1")
    c_num = c4.number_input("参考判例库条目数量", 1, 20, 2, key="c1")
    
    if 'current_user_input' not in st.session_state: st.session_state.current_user_input = ""
    user_input = st.text_area("请输入茶评描述:", value=st.session_state.current_user_input, height=150, key="ui")
    st.session_state.current_user_input = user_input
    
    if 'last_scores' not in st.session_state: 
        st.session_state.last_scores = None
        st.session_state.last_master_comment = ""
    
    # 用于生成动态key，确保每次新评分时校准输入框显示新内容
    if 'score_version' not in st.session_state:
        st.session_state.score_version = 0
        
    if st.button("开始评分", type="primary", use_container_width=True):
        if not user_input:
            st.warning("请输入内容")
        else:
            with st.spinner(f"正在使用 {model_id} 品鉴."):
                user_input = llm_normalize_user_input(user_input, client_d)
                st.session_state.current_user_input = user_input
    
                scores, kb_h, case_h = run_scoring(
                    user_input,
                    st.session_state.kb,
                    st.session_state.cases,
                    st.session_state.prompt_config,
                    embedder,
                    client,
                    "Qwen2.5-7B-Instruct",
                    r_num,
                    c_num
                )
    
                # ✅ 只在这里保存：此时 case_h 一定已定义
                st.session_state.last_case_hits = case_h or []
                st.session_state.last_kb_hits = kb_h or []
    
                if scores:
                    st.session_state.last_scores = scores
                    st.session_state.last_master_comment = scores.get("master_comment", "")
                    st.session_state.score_version += 1
                    st.rerun()


        if st.session_state.last_scores:
            s = st.session_state.last_scores["scores"]
            mc = st.session_state.last_master_comment
            st.markdown(f'<div class="master-comment"><b>👵 宗师总评：</b><br>{mc}</div>', unsafe_allow_html=True)
            

    # ✅ Debug: 展示本次命中的判例（rerun 后仍可见）
    st.caption(f"case_data 条数 = {len(st.session_state.cases[1])} | case_index.ntotal = {st.session_state.cases[0].ntotal}")
    case_h = st.session_state.get("last_case_hits", [])
    st.subheader("🔍 Debug: 命中的判例（Top-K）")
    if case_h:
        for j, c in enumerate(case_h[:c_num], start=1):
            st.markdown(f"**#{j}** {c.get('text','')[:80]}...")
            st.caption(" | ".join([f"{k}:{v.get('score')}" for k,v in (c.get('scores') or {}).items()]))
    else:
        st.warning("Debug: 未命中任何判例（case_h 为空）")
    s = (st.session_state.last_scores or {}).get("scores", {}) or {}
    mc = st.session_state.get("last_master_comment", "")
    factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
    left_col, right_col = st.columns([35, 65]) 
    with left_col:
        st.subheader("📊 风味形态")
        st.pyplot(plot_flavor_shape(st.session_state.last_scores), use_container_width=True)
    with right_col:
        cols = st.columns(2)
        factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
        for i, f in enumerate(factors):
            if f in s:
                d = s[f]
                with cols[i%2]:
                    st.markdown(f"""<div class="factor-card"><div class="score-header"><span>{f}</span><span>{d['score']}/9</span></div><div>{d['comment']}</div><div class="advice-tag">💡 {d.get('suggestion','')}</div></div>""", unsafe_allow_html=True)
    
    st.subheader("🛠️ 评分校准与修正")
    v = st.session_state.score_version  # 获取当前版本号
    mc = st.session_state.get("last_master_comment", "")
    cal_master = st.text_area("校准总评", mc, key=f"cal_master_{v}")
    cal_scores = {}
    st.write("分项调整")
    active_factors = [f for f in factors if f in s]
    grid_cols = st.columns(3) 
    for i, f in enumerate(active_factors):
        with grid_cols[i % 3]:
            with st.container(border=True):
                t_col, s_col = st.columns([1, 1])
                with t_col:
                    st.markdown(f"<div style='padding-top: 5px;'><b>📌 {f}</b></div>", unsafe_allow_html=True)
                with s_col:
                    new_score = st.number_input("分数", 0, 9, int(s[f]['score']), 1, key=f"s_{f}_{v}", label_visibility="collapsed")
                cal_scores[f] = {
                    "score": new_score,
                    "comment": st.text_area(f"评语", s[f]['comment'], key=f"c_{f}_{v}", height=80, placeholder="评语"),
                    "suggestion": st.text_area(f"建议", s[f].get('suggestion',''), key=f"sg_{f}_{v}", height=68, placeholder="建议")
                }
    if st.button("💾 保存校准评分", type="primary"):
        # A. 构造专家数据包
        expert_package = {"scores": cal_scores, "master_comment": cal_master}
        # B. 构造 AI 数据包 (确保 st.session_state.last_scores 存在)
        ai_package = st.session_state.last_scores

        with st.spinner("同步数据到云端记忆模块..."):
            # 1. 存入判例库 (原有逻辑)
            nc_text = st.session_state.get("current_user_input", user_input)
            nc = {"text": nc_text, "scores": cal_scores, "tags": "交互-校准", "master_comment": cal_master, "created_at": time.strftime("%Y-%m-%d")}
            st.session_state.cases[1].append(nc)
            
            # ✅ embedding 与 nc["text"] 完全一致
            v = embedder.encode([nc_text]).astype("float32")
            faiss.normalize_L2(v)
            st.session_state.cases[0].add(v)
            ResourceManager.save(st.session_state.cases[0], st.session_state.cases[1], PATHS.case_index, PATHS.case_data, is_json=True)
            GithubSync.sync_cases(st.session_state.cases[1])
            
            # 2. 存入评测日志 (新增逻辑：LLM-as-a-judge 的原料)
            EvaluationLogger.log_evaluation(
                text= st.session_state.get("current_user_input", user_input), 
                model_output=ai_package, 
                expert_output=expert_package
            )
        
        st.success("校准已存入判例库，误差数据已归档！")
        time.sleep(1)
        st.rerun()

# --- Tab 2: 批量评分 ---
with tab2:
    f = st.file_uploader("上传文件 (.txt/.docx)")
    c1, c2, c3, c4, c5 = st.columns([1, 3, 1, 3, 1])
    r_n = c2.number_input("参考知识库条目数量", 1, 20, 3, key="rb")
    c_n = c4.number_input("参考判例库条目数量", 1, 20, 2, key="cb")
    if f and st.button("批量处理"):
        lines = [l.strip() for l in parse_file(f).split('\n') if len(l)>10]
        res, bar = [], st.progress(0)
        for i, l in enumerate(lines):
            l = llm_normalize_user_input(l, client_d)
            s, _, _ = run_scoring(l, st.session_state.kb, st.session_state.cases, st.session_state.prompt_config, embedder, client, "Qwen2.5-7B-Instruct", r_n, c_n)
            res.append({"id":i+1, "text":l, "scores":s})
            bar.progress((i+1)/len(lines))
        st.success("完成")
        st.download_button("下载Word", create_word_report(res), "report.docx")
    
    # --- Tab 3: RAG ---
with tab3:
    st.subheader("📚 知识库 (RAG)")
    st.caption("上传PDF/文档以增强模型回答的准确性。文件将同步到云端。")
    colu1, colu2 = st.columns([7,3])
    with colu1:
        # ===== 显示GitHub上的RAG文件列表 =====
        st.markdown("**📁 云端上的RAG文件：**")
        
        # 获取GitHub上的文件列表
        if 'github_rag_files' not in st.session_state:
            st.session_state.github_rag_files = []
        
        col_refresh, col_spacer = st.columns([1, 3])
        with col_refresh:
            if st.button("🔄 刷新列表", key="refresh_rag_list"):
                with st.spinner("正在获取文件列表..."):
                    st.session_state.github_rag_files = GithubSync.list_rag_files()
                st.rerun()
        
        github_files = st.session_state.github_rag_files
    if not github_files:
        # 首次加载时尝试获取
        github_files = GithubSync.list_rag_files()
        st.session_state.github_rag_files = github_files
    
    if github_files:
        st.info(f"共 {len(github_files)} 个文件")
        
        # 用于追踪需要删除的文件
        if 'rag_files_to_delete' not in st.session_state:
            st.session_state.rag_files_to_delete = set()
        
        # 显示文件列表，每个文件带删除按钮
        for fname in github_files:
            file_col, del_col = st.columns([5, 1])
            with file_col:
                if fname in st.session_state.rag_files_to_delete:
                    st.markdown(f"~~📄 {fname}~~ *(待删除)*")
                else:
                    st.markdown(f"📄 {fname}")
            with del_col:
                if fname not in st.session_state.rag_files_to_delete:
                    if st.button("🗑️", key=f"del_rag_{fname}", help=f"删除 {fname}"):
                        st.session_state.rag_files_to_delete.add(fname)
                        st.rerun()
                else:
                    if st.button("↩️", key=f"undo_rag_{fname}", help="撤销删除"):
                        st.session_state.rag_files_to_delete.discard(fname)
                        st.rerun()
        
        # 如果有待删除的文件，显示确认按钮
        if st.session_state.rag_files_to_delete:
            st.warning(f"⚠️ 将删除 {len(st.session_state.rag_files_to_delete)} 个文件")
            del_col1, del_col2 = st.columns(2)
            with del_col1:
                if st.button("✅ 确认删除", type="primary", key="confirm_del_rag"):
                    with st.spinner("正在删除文件..."):
                        deleted = []
                        for fname in st.session_state.rag_files_to_delete:
                            if GithubSync.delete_rag_file(fname):
                                deleted.append(fname)
                        
                        # 更新session state
                        st.session_state.github_rag_files = [f for f in github_files if f not in deleted]
                        
                        # 更新本地知识库文件列表
                        current_kb_files = st.session_state.get('kb_files', [])
                        st.session_state.kb_files = [f for f in current_kb_files if f not in deleted]
                        ResourceManager.save_kb_files(st.session_state.kb_files)
                        
                        st.session_state.rag_files_to_delete = set()
                        st.success(f"✅ 已删除 {len(deleted)} 个文件")
                        
                        # 提示需要重建知识库
                        st.info("💡 文件已从云端删除。如需更新本地知识库，请点击下方的'重建本地知识库'按钮。")
                        time.sleep(1)
                        st.rerun()
            with del_col2:
                if st.button("❌ 取消", key="cancel_del_rag"):
                    st.session_state.rag_files_to_delete = set()
                    st.rerun()
    else:
        st.caption("暂无RAG文件")
    
    st.markdown("---")
    
    # ===== 上传新文件（添加模式） =====
    st.markdown("**➕ 添加新文件：**")
    up = st.file_uploader("选择文件", accept_multiple_files=True, key="kb_uploader", 
                        type=['pdf', 'txt', 'docx'])
    
    if up and st.button("📤 添加到知识库", type="primary"):
        # 检查是否有重名文件
        new_names = [u.name for u in up]
        existing_names = st.session_state.get('github_rag_files', [])
        duplicate_names = set(new_names) & set(existing_names)
        
        if duplicate_names:
            st.warning(f"⚠️ 以下文件已存在，将被覆盖：{', '.join(duplicate_names)}")
        
        with st.spinner("正在处理文件..."):
            # 1. 解析文件内容
            raw = "".join([parse_file(u) for u in up])
            
            if not raw.strip():
                st.error("❌ 无法从上传的文件中提取有效文本")
            else:
                # 2. 上传到GitHub
                with st.spinner("上传到GitHub..."):
                    success, uploaded_names = GithubSync.add_rag_files(up, "tea_data/RAG")
                
                if success:
                    # 3. 更新本地文件列表
                    current_kb_files = st.session_state.get('kb_files', [])
                    # 合并文件列表（去重）
                    all_files = list(set(current_kb_files + uploaded_names))
                    st.session_state.kb_files = all_files
                    st.session_state.github_rag_files = list(set(existing_names + uploaded_names))
                    ResourceManager.save_kb_files(all_files)
                    
                    st.success(f"✅ 已上传 {len(uploaded_names)} 个文件到GitHub")
                    st.info("💡 请点击下方的'重建本地知识库'按钮以更新向量索引。")
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error("❌ 上传到GitHub失败")

# ===== 重建本地知识库按钮 =====
with colu2:
    st.markdown("**🔧 知识库维护：**")
    local_kb_count = len(st.session_state.kb[1])
    st.caption(f"网页端知识库：{local_kb_count} 个片段")
    
    # 每个文件换行显示
    if kb_files:
        st.markdown("**网页端知识库文件:**")
        for fname in kb_files:
            st.markdown(f"- 📄 {fname}")
    else:
        st.markdown("**网页端知识库文件:** 无") 
    st.markdown("---")
    st.markdown("云端数据与网页数据不统一？")
    if st.button("🔄 从云端加载知识库", use_container_width=True, type="primary"):
        with st.spinner("正在从云端拉取并重建知识库..."):
            success, msg = load_rag_from_github(aliyun_key)
            if success:
                st.success(msg)
                # 更新GitHub文件列表
                st.session_state.github_rag_files = GithubSync.list_rag_files()
            else:
                st.error(msg)
        time.sleep(1)
        st.rerun()



with tab4:
    MANAGER_URL = "http://117.50.89.74:8001"
    c1, c2 = st.columns([5, 5])
    
    with c1:
        st.subheader("📕 判例库 (CASE)")        
        if st.button("📋 展示当前判例", use_container_width=True):
            show_cases_dialog(embedder)
        
        # 检查是否需要打开编辑弹窗
        if st.session_state.get('editing_case_idx') is not None:
            edit_case_dialog(st.session_state.editing_case_idx, embedder)
        
        with st.expander("➕ 手动添加精细判例"):
            with st.form("case_form"):
                f_txt = st.text_area("判例描述", height=80)
                f_tag = st.text_input("标签", "人工录入")
                st.markdown("**因子评分详情**")
                fc1, fc2 = st.columns(2)
                factors = ["优雅性", "辨识度", "协调性", "饱和度", "持久性", "苦涩度"]
                input_scores = {}
                for i, f in enumerate(factors):
                    with (fc1 if i%2==0 else fc2):
                        val = st.number_input(f"{f}分数", 0,9,7, key=f"s_{i}")
                        cmt = st.text_input(f"{f}评语", key=f"c_{i}")
                        sug = st.text_input(f"{f}建议", key=f"a_{i}")
                        input_scores[f] = {"score": val, "comment": cmt, "suggestion": sug}
                
                if st.form_submit_button("保存判例并同步"):
                    new_c = {"text": f_txt, "tags": f_tag, "scores": input_scores, "created_at": time.strftime("%Y-%m-%d")}
                    st.session_state.cases[1].append(new_c)
                    vec = embedder.encode([f_txt]).astype("float32")
                    faiss.normalize_L2(vec)
                    st.session_state.cases[0].add(vec)
                    ResourceManager.save(st.session_state.cases[0], st.session_state.cases[1], PATHS.case_index, PATHS.case_data, is_json=True)
                    
                    # 同步到GitHub
                    with st.spinner("同步到GitHub..."):
                        GithubSync.sync_cases(st.session_state.cases[1])
                    
                    st.success("已保存并同步！")
                    time.sleep(1); st.rerun()
    
    # --- 右侧：微调控制 ---
    with c2:
        st.subheader("🚀 模型微调 (LoRA)")
        
        server_status = "unknown"
        try:
            resp = requests.get(f"{MANAGER_URL}/status", timeout=2)
            if resp.status_code == 200:
                status_data = resp.json()
                if status_data.get("vllm_status") == "running":
                    server_status = "idle"
                else:
                    server_status = "training"
            else:
                server_status = "error"
        except:
            server_status = "offline"
        
        if server_status == "idle":
            st.success("🟢 服务器就绪 (正在进行推理服务)")
        elif server_status == "training":
            st.warning("🟠 正在微调训练中... (推理服务暂停)")
            st.markdown("⚠️ **注意：** 此时无法进行评分交互，请耐心等待训练完成。")
        elif server_status == "offline":
            st.error("🔴 无法连接到 GPU 服务器 (请联系管理员)")
    
        st.markdown("#### 1. 数据准备")
        
        if PATHS.training_file.exists():
            with open(PATHS.training_file, "r", encoding="utf-8") as f:
                lines = f.readlines()
            data_count = len(lines)
        else:
            data_count = 0
            
        st.info(f"当前微调数据：**{data_count} 条** | 判例库：**{len(st.session_state.cases[1])} 条**")
        
        # ===== 修改：覆盖逻辑 =====
        if st.button("🔄 将当前所有判例转为微调数据（覆盖）"):
            cnt = ResourceManager.overwrite_finetune(
                st.session_state.cases[1],
                st.session_state.prompt_config.get('system_template',''), 
                st.session_state.prompt_config.get('user_template','')
            )
            st.success(f"已覆盖写入 {cnt} 条微调数据！")
            time.sleep(1); st.rerun()
    
        st.markdown("#### 2. 启动训练")
        st.caption("点击下方按钮将把数据上传至 GPU 服务器并开始训练。训练期间服务将中断约 2-5 分钟。")
    
        btn_disabled = (server_status != "idle") or (data_count == 0)
        
        if st.button("🔥 开始微调 (Start LoRA)", type="primary", disabled=btn_disabled):
            if not PATHS.training_file.exists():
                st.error("找不到训练数据文件！")
            else:
                try:
                    with open(PATHS.training_file, "rb") as f:
                        with st.spinner("正在上传数据并启动训练任务..."):
                            files = {'file': ('tea_feedback.jsonl', f, 'application/json')}
                            r = requests.post(f"{MANAGER_URL}/upload_and_train", files=files, timeout=100)
                            
                        if r.status_code == 200:
                            st.balloons()
                            st.success(f"✅ 任务已提交！服务器响应: {r.json().get('message')}")
                            st.info("💡 你可以稍后刷新页面查看状态，训练完成后服务会自动恢复。")
                        else:
                            st.error(f"❌ 提交失败: {r.text}")
                except Exception as e:
                    st.error(f"❌ 连接错误: {e}")

# --- Tab 4: Prompt配置 ---
with tab5:
    pc = st.session_state.prompt_config
    st.markdown("系统提示词**可以修改**。完整全面的提示词会让大语言模型返回的更准确结果。")    
    sys_t = st.text_area("系统提示词", pc.get('system_template',''), height=350)
    st.markdown("用户提示词**不可修改**。其保证了发送内容与回答内容的基本结构，因此大语言模型的回答可被准确解析。")
    user_t = st.text_area("用户提示词", pc.get('user_template',''), height=250, disabled=True)
    
    if st.button("💾 保存（永久化同步）", type="primary"):
        if sys_t == pc.get('system_template'):
            st.info("内容没有变化，无需保存。")
        else:
            new_cfg = {"system_template": sys_t, "user_template": user_t}
            
            with st.spinner("正在连接云端仓库并写入数据..."):
                success = GithubSync.push_json(
                    file_path_in_repo="tea_data/prompts.json", 
                    data_dict=new_cfg,
                    commit_msg="Update prompts.json from App"
                )
            
            if success:
                st.success("✅ 成功写入云端！")
                st.session_state.prompt_config = new_cfg
                with open(PATHS.prompt_config_file, 'w', encoding='utf-8') as f:
                    json.dump(new_cfg, f, ensure_ascii=False, indent=2)

with tab6:
    st.header("🧠 模型效果量化与误差分析（基于日志）")
    
    logs = EvaluationLogger.load_logs() or []
    logs = [l for l in logs if isinstance(l, dict)]
    
    # 只统计有“专家真值”的样本
    paired = [
        l for l in logs
        if l.get("model_prediction") and l.get("expert_ground_truth")
    ]
    
    total = len(logs)
    paired_n = len(paired)
    st.metric("日志总数", total)
    st.metric("可评估样本（有专家真值）", paired_n)
    
    if paired_n == 0:
        st.info("暂无可量化的样本：需要先在交互评分里保存专家校准（expert_ground_truth）。")
    else:
        # --- 计算指标 ---
        per_factor_abs = {}   # factor -> list[abs_err]
        per_factor_signed = {}# factor -> list[signed_err] (model - expert)
        case_errors = []      # (total_abs_err, log_dict)
    
        for l in paired:
            m_scores = (l.get("model_prediction") or {}).get("scores", {}) or {}
            e_scores = (l.get("expert_ground_truth") or {}).get("scores", {}) or {}
    
            abs_list = []
            for factor, m_item in m_scores.items():
                e_item = e_scores.get(factor)
                if not isinstance(m_item, dict) or not isinstance(e_item, dict):
                    continue
                ms = m_item.get("score")
                es = e_item.get("score")
                if not isinstance(ms, (int, float)) or not isinstance(es, (int, float)):
                    continue
    
                signed = ms - es
                abs_err = abs(signed)
    
                per_factor_abs.setdefault(factor, []).append(abs_err)
                per_factor_signed.setdefault(factor, []).append(signed)
                abs_list.append(abs_err)
    
            # 该条样本的平均绝对误差（跨维度）
            if abs_list:
                case_errors.append((sum(abs_list) / len(abs_list), l))
    
        # 总体 MAE（跨所有维度的平均绝对误差）
        all_abs = [x for xs in per_factor_abs.values() for x in xs]
        overall_mae = sum(all_abs) / len(all_abs) if all_abs else 0.0
    
        # 方向性偏差：平均 (model - expert)
        all_signed = [x for xs in per_factor_signed.values() for x in xs]
        overall_bias = sum(all_signed) / len(all_signed) if all_signed else 0.0
    
        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("总体 MAE（分）", f"{overall_mae:.3f}")
        with c2:
            st.metric("总体偏差（model-expert）", f"{overall_bias:+.3f}")
        with c3:
            st.metric("校准覆盖率", f"{paired_n/total:.1%}" if total else "0%")
    
        st.divider()
    
        # --- 每维度指标 ---
        st.subheader("📊 各维度误差（MAE）与偏差方向")
        rows = []
        for factor in sorted(per_factor_abs.keys()):
            abs_errs = per_factor_abs[factor]
            signed_errs = per_factor_signed.get(factor, [])
            mae = sum(abs_errs) / len(abs_errs) if abs_errs else 0.0
            bias = sum(signed_errs) / len(signed_errs) if signed_errs else 0.0
            rows.append((factor, mae, bias, len(abs_errs)))
    
        # 用 st.dataframe 展示（不依赖 pandas）
        st.dataframe(
            [{"factor": f, "mae": round(mae, 3), "bias(model-expert)": round(bias, 3), "n": n}
             for (f, mae, bias, n) in rows],
            use_container_width=True
        )
    
        st.divider()
    
        # --- Top-N 误差样本定位 ---
        st.subheader("🔎 误差最大样本 Top-N（用于定位问题）")
        topn = st.slider("Top-N", min_value=3, max_value=30, value=10, step=1)
    
        case_errors.sort(key=lambda x: x[0], reverse=True)
        for rank, (err, l) in enumerate(case_errors[:topn], start=1):
            ts = l.get("timestamp", "unknown")
            txt = (l.get("input_text") or "")
            title = f"#{rank} | 平均误差={err:.3f} | {ts} | 输入: {txt[:20]}..."
            with st.expander(title):
                col1, col2 = st.columns(2)
                with col1:
                    st.caption("🤖 模型输出")
                    st.json(l.get("model_prediction", {}))
                with col2:
                    st.caption("👨‍🏫 专家真值")
                    st.json(l.get("expert_ground_truth", {}))
    
                # 可选：一键让 AI 写“差异原因分析”
                if not l.get("analysis"):
                    if st.button("⚖️ 让 AI 分析差异原因（写入日志）", key=f"judge_{l.get('id','noid')}"):
                        with st.spinner("AI 正在生成差异原因分析..."):
                            EvaluationLogger.run_judge(l["id"], client_d)  # 你项目里一般叫 client_d
                            st.success("完成，已写入日志 analysis 字段")
                            st.rerun()
                else:
                    st.info(l["analysis"])
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    
    










